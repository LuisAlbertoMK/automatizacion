"""
modules/documentos/escrito.py — Generador de escritos, cartas y documentos con Claude API → .docx.

Migrado de tramites-auto/tramites-bot/docs/escrito.js.
Soporta 11 tipos de documentos: cartas, solicitudes, contratos, quejas, etc.

Uso:
    from src.tramites.documentos import EscritoGenerator
    generador = EscritoGenerator()
    resultado = generador.generar_interactivo()
"""

import json
import os
from datetime import datetime
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

from src.utils.claude import call_claude

# ── Estilos ──────────────────────────────────────────────────────────────────
COLOR_TEXT = RGBColor(0x33, 0x33, 0x33)
COLOR_MUTED = RGBColor(0x88, 0x88, 0x88)

OUTPUT_DIR = Path(__file__).parent.parent.parent.parent / "output"


TIPOS = {
    "1": "carta_presentacion",
    "2": "carta_recomendacion",
    "3": "solicitud_empleo",
    "4": "poder_notarial_simple",
    "5": "carta_responsiva",
    "6": "carta_de_no_adeudo",
    "7": "queja_formal",
    "8": "demanda_conciliacion",
    "9": "contrato_prestacion_servicios",
    "10": "control_confianza_redaccion",
    "11": "escrito_libre",
}

DESCRIPCIONES = {
    "carta_presentacion": "carta de presentación profesional",
    "carta_recomendacion": "carta de recomendación laboral",
    "solicitud_empleo": "solicitud formal de empleo",
    "poder_notarial_simple": "poder notarial simple (no requiere notario)",
    "carta_responsiva": "carta responsiva con cláusulas claras",
    "carta_de_no_adeudo": "carta de no adeudo / finiquito",
    "queja_formal": "escrito de queja formal ante autoridad",
    "demanda_conciliacion": "demanda ante centro de conciliación laboral",
    "contrato_prestacion_servicios": "contrato de prestación de servicios profesionales",
    "control_confianza_redaccion": "redacción de respuestas para cuestionario de control de confianza",
    "escrito_libre": "escrito libre según indicaciones del usuario",
}


def _parrafo(doc, texto, size=11, bold=False, color=COLOR_TEXT,
             alignment=WD_ALIGN_PARAGRAPH.JUSTIFY, space_after=160):
    """Agrega un párrafo con formato consistente."""
    p = doc.add_paragraph()
    run = p.add_run(texto)
    run.font.size = Pt(size)
    run.font.name = "Arial"
    run.font.color.rgb = color
    run.bold = bold
    p.alignment = alignment
    p.paragraph_format.space_after = Pt(space_after / 20)  # twips → pt
    return p


class EscritoGenerator:
    """
    Generador de escritos, cartas y documentos legales con Claude API.
    """

    def _generar_escrito_con_claude(self, tipo: str, datos: dict) -> dict:
        """Llama a Claude para redactar el contenido del documento."""
        descripcion = DESCRIPCIONES.get(tipo, tipo)

        prompt = f"""Redacta un {descripcion} en español, formal y profesional.
Responde SOLO con JSON válido sin markdown.

Estructura:
{{
  "titulo": "...",
  "lugar_fecha": "Ciudad, DD de mes de AAAA",
  "destinatario": "...",
  "asunto": "...",
  "cuerpo": ["párrafo 1", "párrafo 2", ...],
  "cierre": "...",
  "firmante": "...",
  "notas_legales": "advertencia breve si aplica (o vacío)"
}}

Datos disponibles:
{json.dumps(datos, indent=2, ensure_ascii=False)}

Tipo de documento: {tipo}
Instrucciones adicionales: {datos.get('instrucciones', 'ninguna')}"""

        return call_claude([{"role": "user", "content": prompt}], max_tokens=2000)

    def _construir_docx(self, doc_data: dict) -> Document:
        """Construye el documento Word con el escrito."""
        doc = Document()

        # Configurar fuente default
        style = doc.styles["Normal"]
        style.font.name = "Arial"
        style.font.size = Pt(11)
        style.font.color.rgb = COLOR_TEXT

        # Márgenes (1 pulgada = 1440 twips, como el original)
        for section in doc.sections:
            section.top_margin = Inches(1)
            section.bottom_margin = Inches(1)
            section.left_margin = Inches(1)
            section.right_margin = Inches(1)

        # ── Título ──
        titulo = doc.add_paragraph()
        titulo.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = titulo.add_run(doc_data.get("titulo", "").upper())
        run.bold = True
        run.font.size = Pt(13)
        run.font.name = "Arial"
        titulo.paragraph_format.space_after = Pt(20)

        # ── Lugar y fecha ──
        lugar_fecha = doc_data.get("lugar_fecha", "")
        if lugar_fecha:
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            run = p.add_run(lugar_fecha)
            run.italic = True
            run.font.size = Pt(10)
            run.font.name = "Arial"
            p.paragraph_format.space_after = Pt(16)

        # ── Destinatario ──
        destinatario = doc_data.get("destinatario", "")
        if destinatario:
            _parrafo(doc, destinatario, bold=True, space_after=160)

        # ── Asunto ──
        asunto = doc_data.get("asunto", "")
        if asunto:
            p = doc.add_paragraph()
            run = p.add_run("ASUNTO: ")
            run.bold = True
            run.font.size = Pt(11)
            run.font.name = "Arial"
            run = p.add_run(asunto)
            run.font.size = Pt(11)
            run.font.name = "Arial"
            p.paragraph_format.space_after = Pt(16)

        # ── Cuerpo ──
        for parrafo_texto in doc_data.get("cuerpo", []):
            _parrafo(doc, parrafo_texto, space_after=160)

        # ── Cierre ──
        doc.add_paragraph().paragraph_format.space_after = Pt(20)

        cierre = doc_data.get("cierre", "")
        if cierre:
            _parrafo(doc, cierre, space_after=40)

        # ── Firma ──
        firmante = doc_data.get("firmante", "")
        if firmante:
            p = doc.add_paragraph()
            pPr = p._p.get_or_add_pPr()
            pBdr = pPr.makeelement(qn("w:pBdr"), {})
            top = pBdr.makeelement(qn("w:top"), {
                qn("w:val"): "single",
                qn("w:sz"): "4",
                qn("w:color"): "333333",
                qn("w:space"): "3",
            })
            pBdr.append(top)
            pPr.append(pBdr)

            run = p.add_run(firmante)
            run.bold = True
            run.font.size = Pt(11)
            run.font.name = "Arial"
            p.paragraph_format.space_before = Pt(40)

        # ── Notas legales ──
        notas = doc_data.get("notas_legales", "")
        if notas:
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(20)
            run = p.add_run(f"⚠️ {notas}")
            run.font.size = Pt(8)
            run.font.name = "Arial"
            run.font.color.rgb = COLOR_MUTED
            run.italic = True

        return doc

    def _preguntar(self, mensaje: str) -> str:
        """Pide un dato al usuario."""
        return input(f"  {mensaje}: ").strip()

    def generar_interactivo(self) -> dict:
        """
        Modo interactivo: menú de tipos, pide datos, genera documento.

        Returns:
            dict con ruta del archivo generado
        """
        print("\n" + "=" * 60)
        print("  📄 GENERADOR DE DOCUMENTOS Y ESCRITOS")
        print("=" * 60)
        print("\n  Tipos disponibles:")
        for k, v in TIPOS.items():
            print(f"    {k}. {v.replace('_', ' ')}")

        tipo_num = self._preguntar("\n  Elige tipo (número)")
        tipo = TIPOS.get(tipo_num, "escrito_libre")
        print(f"  → {tipo.replace('_', ' ')}")

        datos = {}
        datos["tipo"] = tipo
        datos["instrucciones"] = self._preguntar(
            "Describe qué necesitas (más detalles = mejor resultado)"
        )
        datos["destinatario"] = self._preguntar(
            "Dirigido a (nombre / empresa / autoridad)"
        )
        datos["fecha"] = self._preguntar("Fecha (ej: 27 de marzo de 2026)")
        datos["nombre"] = self._preguntar("Tu nombre completo")

        print("\n  🤖 Redactando con IA...")
        try:
            escrito = self._generar_escrito_con_claude(tipo, datos)
        except Exception as e:
            print(f"  ❌ Error al generar documento: {e}")
            return {"status": "error", "error": str(e)}

        # Completar datos faltantes
        escrito["lugar_fecha"] = f"{datos.get('ciudad', 'México')}, {datos['fecha']}"
        escrito["firmante"] = datos["nombre"]

        print("  📝 Construyendo documento...")
        doc = self._construir_docx(escrito)

        # Guardar
        os.makedirs(OUTPUT_DIR, exist_ok=True)
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        nombre_archivo = f"{tipo}-{timestamp}.docx"
        ruta = OUTPUT_DIR / nombre_archivo
        doc.save(str(ruta))
        print(f"  ✅ Documento generado: {ruta}")

        return {
            "status": "ok",
            "archivo": str(ruta),
            "tipo": tipo,
        }
