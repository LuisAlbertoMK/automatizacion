"""
modules/documentos/cv.py — Generador de CV profesional con Claude API → .docx.

Migrado de tramites-auto/tramites-bot/docs/cv.js.
Usa python-docx para construir el documento Word y Claude API para
generar el contenido estructurado del CV.

Uso:
    from modules.documentos import CVGenerator
    generador = CVGenerator()
    resultado = generador.generar_interactivo()
"""

import json
import os
import re
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

from utils.claude import call_claude

# ── Estilos ──────────────────────────────────────────────────────────────────
COLOR_PRIMARY = RGBColor(0x1F, 0x38, 0x64)   # Azul oscuro
COLOR_ACCENT = RGBColor(0x2E, 0x75, 0xB6)    # Azul medio
COLOR_TEXT = RGBColor(0x33, 0x33, 0x33)       # Gris casi negro
COLOR_MUTED = RGBColor(0x66, 0x66, 0x66)      # Gris


OUTPUT_DIR = Path(__file__).parent.parent.parent.parent / "output"


def _add_line(doc, color=COLOR_ACCENT):
    """Agrega línea horizontal decorativa."""
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    pBdr = pPr.makeelement(qn("w:pBdr"), {})
    bottom = pBdr.makeelement(qn("w:bottom"), {
        qn("w:val"): "single",
        qn("w:sz"): "6",
        qn("w:color"): f"{color[0]:02X}{color[1]:02X}{color[2]:02X}",
        qn("w:space"): "1",
    })
    pBdr.append(bottom)
    pPr.append(pBdr)
    p.paragraph_format.space_after = Pt(6)
    return p


def _seccion(doc, titulo):
    """Agrega título de sección con línea decorativa."""
    _add_line(doc)
    p = doc.add_paragraph()
    run = p.add_run(titulo.upper())
    run.bold = True
    run.font.size = Pt(12)
    run.font.color.rgb = COLOR_ACCENT
    run.font.name = "Arial"
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after = Pt(6)


def _bullet(doc, texto):
    """Agrega un párrafo con viñeta."""
    p = doc.add_paragraph(style="List Bullet")
    run = p.add_run(texto)
    run.font.size = Pt(10)
    run.font.color.rgb = COLOR_TEXT
    run.font.name = "Arial"
    p.paragraph_format.space_after = Pt(2)


def _parrafo(doc, texto, size=10, bold=False, color=COLOR_TEXT, alignment=None):
    """Agrega un párrafo simple."""
    p = doc.add_paragraph()
    run = p.add_run(texto)
    run.font.size = Pt(size)
    run.font.name = "Arial"
    run.font.color.rgb = color
    run.bold = bold
    if alignment:
        p.alignment = alignment
    return p


class CVGenerator:
    """
    Generador de CV profesional usando Claude API.
    
    El flujo es:
      1. Recibe datos del usuario (interactivo o directo)
      2. Claude genera contenido estructurado del CV
      3. Se construye un documento .docx con python-docx
    """

    def __init__(self, api_key: str | None = None):
        self.api_key = api_key

    # ── Prompt a Claude ─────────────────────────────────────────────────────

    def _generar_cv_con_claude(self, datos: dict) -> dict:
        """Llama a Claude para generar el contenido del CV."""
        prompt = f"""Genera un CV profesional en español para la siguiente persona.
Responde SOLO con JSON válido sin markdown ni backticks.

Estructura requerida:
{{
  "nombre_completo": "",
  "titulo_profesional": "",
  "resumen": "3 oraciones máximo",
  "experiencia": [
    {{ "empresa": "", "puesto": "", "periodo": "", "logros": [""] }}
  ],
  "educacion": [
    {{ "institucion": "", "grado": "", "año": "" }}
  ],
  "habilidades": [""],
  "idiomas": [""],
  "contacto": {{ "telefono": "", "email": "", "ciudad": "" }}
}}

Datos del cliente:
{json.dumps(datos, indent=2, ensure_ascii=False)}

Completa los campos faltantes de forma profesional según el puesto objetivo.
Si falta información, usa placeholders descriptivos como "[Agregar empresa]"."""

        return call_claude([{"role": "user", "content": prompt}], max_tokens=1500)

    # ── Construcción del .docx ──────────────────────────────────────────────

    def _construir_docx(self, cv: dict) -> Document:
        """Construye el documento Word con el CV."""
        doc = Document()

        # Configurar fuente default
        style = doc.styles["Normal"]
        style.font.name = "Arial"
        style.font.size = Pt(10)
        style.font.color.rgb = COLOR_TEXT

        # Márgenes
        for section in doc.sections:
            section.top_margin = Inches(0.79)
            section.bottom_margin = Inches(0.79)
            section.left_margin = Inches(0.79)
            section.right_margin = Inches(0.79)

        # ── Encabezado ──
        nombre = doc.add_paragraph()
        nombre.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = nombre.add_run(cv.get("nombre_completo", ""))
        run.bold = True
        run.font.size = Pt(18)
        run.font.color.rgb = COLOR_PRIMARY
        run.font.name = "Arial"
        nombre.paragraph_format.space_after = Pt(4)

        titulo = doc.add_paragraph()
        titulo.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = titulo.add_run(cv.get("titulo_profesional", ""))
        run.italic = True
        run.font.size = Pt(12)
        run.font.color.rgb = COLOR_ACCENT
        run.font.name = "Arial"
        titulo.paragraph_format.space_after = Pt(4)

        contacto = cv.get("contacto", {})
        datos_contacto = "  |  ".join(
            filter(None, [
                contacto.get("telefono", ""),
                contacto.get("email", ""),
                contacto.get("ciudad", ""),
            ])
        )
        if datos_contacto:
            _parrafo(doc, datos_contacto, size=9, color=COLOR_MUTED,
                     alignment=WD_ALIGN_PARAGRAPH.CENTER)
            doc.paragraphs[-1].paragraph_format.space_after = Pt(12)

        # ── Perfil Profesional ──
        _seccion(doc, "Perfil Profesional")
        _parrafo(doc, cv.get("resumen", ""))
        doc.paragraphs[-1].paragraph_format.space_after = Pt(8)

        # ── Experiencia Laboral ──
        _seccion(doc, "Experiencia Laboral")
        for exp in cv.get("experiencia", []):
            p = doc.add_paragraph()
            run = p.add_run(exp.get("puesto", ""))
            run.bold = True
            run.font.size = Pt(11)
            run.font.name = "Arial"
            run.font.color.rgb = COLOR_TEXT

            empresa = exp.get("empresa", "")
            if empresa:
                run = p.add_run(f"  —  {empresa}")
                run.font.size = Pt(10)
                run.font.name = "Arial"
                run.font.color.rgb = RGBColor(0x44, 0x44, 0x44)

            periodo = exp.get("periodo", "")
            if periodo:
                run = p.add_run(f"  ({periodo})")
                run.font.size = Pt(9)
                run.font.name = "Arial"
                run.font.color.rgb = COLOR_MUTED
                run.italic = True

            p.paragraph_format.space_before = Pt(8)
            p.paragraph_format.space_after = Pt(4)

            for logro in exp.get("logros", []):
                _bullet(doc, logro)

        # ── Educación ──
        _seccion(doc, "Educación")
        for ed in cv.get("educacion", []):
            p = doc.add_paragraph()
            run = p.add_run(ed.get("grado", ""))
            run.bold = True
            run.font.size = Pt(10)
            run.font.name = "Arial"

            institucion = ed.get("institucion", "")
            if institucion:
                run = p.add_run(f"  —  {institucion}")
                run.font.size = Pt(10)
                run.font.name = "Arial"

            anio = ed.get("año", "")
            if anio:
                run = p.add_run(f"  ({anio})")
                run.font.size = Pt(9)
                run.font.name = "Arial"
                run.font.color.rgb = COLOR_MUTED
                run.italic = True

            p.paragraph_format.space_after = Pt(4)

        # ── Habilidades ──
        _seccion(doc, "Habilidades")
        habilidades = cv.get("habilidades", [])
        if habilidades:
            _parrafo(doc, "  •  ".join(habilidades))
            doc.paragraphs[-1].paragraph_format.space_after = Pt(8)

        # ── Idiomas ──
        _seccion(doc, "Idiomas")
        idiomas = cv.get("idiomas", [])
        if idiomas:
            _parrafo(doc, "  •  ".join(idiomas))

        return doc

    # ── Flujo interactivo ────────────────────────────────────────────────────

    def _preguntar(self, mensaje: str, default: str = "") -> str:
        """Pide un dato al usuario."""
        hint = f" [{default}]" if default else ""
        val = input(f"  {mensaje}{hint}: ").strip()
        return val if val else default

    def generar_interactivo(self) -> dict:
        """
        Modo interactivo: pide datos al usuario, genera CV y guarda .docx.

        Returns:
            dict con ruta del archivo generado y datos del CV
        """
        print("\n" + "=" * 60)
        print("  📋 GENERADOR DE CV — responde las siguientes preguntas")
        print("=" * 60)

        datos = {}

        # Cargar perfil desde output/perfiles.json si existe
        perfiles_path = Path(OUTPUT_DIR) / "perfiles.json"
        if perfiles_path.exists():
            try:
                with open(perfiles_path, encoding="utf-8") as f:
                    perfiles = json.load(f)
                if isinstance(perfiles, dict):
                    datos.update(perfiles)
            except (json.JSONDecodeError, OSError):
                pass

        datos["nombre"] = self._preguntar("Nombre completo", datos.get("nombre", ""))
        datos["puesto_objetivo"] = self._preguntar("Puesto al que aplicas")
        datos["anos_experiencia"] = self._preguntar("Años de experiencia")
        datos["ultimo_trabajo"] = self._preguntar("Último trabajo (empresa y puesto)")
        datos["habilidades_principales"] = self._preguntar("3 habilidades principales (separadas por coma)")
        datos["idiomas_input"] = self._preguntar("Idiomas (ej: Español nativo, Inglés B2)")
        datos["ciudad"] = self._preguntar("Ciudad")
        datos["email"] = self._preguntar("Email", datos.get("email", ""))
        datos["telefono"] = self._preguntar("Teléfono", datos.get("telefono", ""))

        print("\n  🤖 Generando CV con IA...")
        try:
            cv = self._generar_cv_con_claude(datos)
        except Exception as e:
            print(f"  ❌ Error al generar CV: {e}")
            return {"status": "error", "error": str(e)}

        print("  📝 Construyendo documento Word...")
        doc = self._construir_docx(cv)

        # Guardar
        sanitized = re.sub(r"\s+", "-", datos["nombre"]).lower()
        nombre_archivo = f"cv-{sanitized}.docx"
        os.makedirs(OUTPUT_DIR, exist_ok=True)
        ruta = OUTPUT_DIR / nombre_archivo
        doc.save(str(ruta))
        print(f"  ✅ CV generado: {ruta}")

        return {
            "status": "ok",
            "archivo": str(ruta),
            "nombre": cv.get("nombre_completo", datos["nombre"]),
        }
