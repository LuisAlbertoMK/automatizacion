"""Tests para src/tramites/documentos/escrito.py — Generador de documentos."""

import json
import os
from datetime import datetime
from pathlib import Path
from unittest.mock import MagicMock, patch

import pytest

from src.tramites.documentos.escrito import (
    OUTPUT_DIR, TIPOS, DESCRIPCIONES, _parrafo, EscritoGenerator,
)

# ── Helpers ──────────────────────────────────────────────────────────────────

_DOC_COMPLETO = {
    "titulo": "Carta de Presentación",
    "lugar_fecha": "CDMX, 27 de marzo de 2026",
    "destinatario": "Lic. Juan Pérez\nDirector de RH\nTech Corp",
    "asunto": "Solicitud de empleo para puesto de Developer",
    "cuerpo": [
        "Por medio de la presente, me permito presentar mi solicitud...",
        "Cuento con 5 años de experiencia en desarrollo de software...",
    ],
    "cierre": "Sin otro particular, quedo a sus órdenes.",
    "firmante": "María García López",
    "notas_legales": "Este documento es una carta de presentación.",
}

_DOC_MINIMO = {
    "titulo": "",
    "lugar_fecha": "",
    "destinatario": "",
    "asunto": "",
    "cuerpo": [],
    "cierre": "",
    "firmante": "",
    "notas_legales": "",
}


@pytest.fixture
def generator():
    return EscritoGenerator()


@pytest.fixture
def tmp_output(tmp_path):
    with patch("src.tramites.documentos.escrito.OUTPUT_DIR", tmp_path):
        yield tmp_path


# ── Constantes ───────────────────────────────────────────────────────────────

class TestConstantes:
    def test_tipos_tiene_11(self):
        assert len(TIPOS) == 11

    def test_descripciones_cubren_tipos(self):
        for v in TIPOS.values():
            assert v in DESCRIPCIONES

    def test_tipo_default_escrito_libre(self):
        assert TIPOS.get("999", "escrito_libre") == "escrito_libre"


# ── _parrafo ─────────────────────────────────────────────────────────────────

class TestParrafo:
    def test_parrafo_basico(self):
        from docx import Document
        from docx.shared import RGBColor

        doc = Document()
        p = _parrafo(doc, "Texto de prueba")
        assert "Texto de prueba" in p.text
        run = p.runs[0]
        assert run.font.size.pt == 11

    def test_parrafo_personalizado(self):
        from docx import Document
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.shared import RGBColor

        doc = Document()
        p = _parrafo(doc, "Bold right", size=14, bold=True,
                     color=RGBColor(0xFF, 0, 0),
                     alignment=WD_ALIGN_PARAGRAPH.RIGHT,
                     space_after=400)
        assert p.alignment == WD_ALIGN_PARAGRAPH.RIGHT
        assert p.runs[0].bold


# ── EscritoGenerator._construir_docx ────────────────────────────────────────

class TestConstruirDocx:
    """Cubre _construir_docx (lines 107-205)."""

    def test_doc_completo(self, generator):
        doc = generator._construir_docx(_DOC_COMPLETO)
        texts = "\n".join(p.text for p in doc.paragraphs)

        assert "CARTA DE PRESENTACIÓN" in texts
        assert "CDMX, 27 de marzo de 2026" in texts
        assert "Lic. Juan Pérez" in texts
        assert "ASUNTO:" in texts
        assert "Solicitud de empleo" in texts
        assert "Por medio de la presente" in texts
        assert "Sin otro particular" in texts
        assert "María García López" in texts
        assert "Este documento es una carta" in texts

    def test_doc_minimo_no_crashea(self, generator):
        """Campos vacíos no deben causar errores."""
        doc = generator._construir_docx(_DOC_MINIMO)
        assert doc is not None
        assert len(doc.paragraphs) >= 2

    def test_sin_lugar_fecha(self, generator):
        doc_data = {**_DOC_COMPLETO, "lugar_fecha": ""}
        doc = generator._construir_docx(doc_data)
        texts = "\n".join(p.text for p in doc.paragraphs)
        # No agrega párrafo de lugar_fecha vacío
        assert "CDMX, 27" not in texts

    def test_sin_destinatario(self, generator):
        doc_data = {**_DOC_COMPLETO, "destinatario": ""}
        doc = generator._construir_docx(doc_data)
        texts = "\n".join(p.text for p in doc.paragraphs)
        assert "ASUNTO:" in texts  # asunto igual se agrega

    def test_sin_asunto(self, generator):
        doc_data = {**_DOC_COMPLETO, "asunto": ""}
        doc = generator._construir_docx(doc_data)
        texts = "\n".join(p.text for p in doc.paragraphs)
        assert "ASUNTO:" not in texts

    def test_sin_cuerpo(self, generator):
        """Cuerpo vacío → no crashea, cierre aparece."""
        doc_data = {**_DOC_COMPLETO, "cuerpo": []}
        doc = generator._construir_docx(doc_data)
        texts = "\n".join(p.text for p in doc.paragraphs)
        assert "Sin otro particular" in texts

    def test_sin_cierre(self, generator):
        doc_data = {**_DOC_COMPLETO, "cierre": ""}
        doc = generator._construir_docx(doc_data)
        texts = "\n".join(p.text for p in doc.paragraphs)
        assert "Sin otro particular" not in texts

    def test_sin_firmante(self, generator):
        doc_data = {**_DOC_COMPLETO, "firmante": ""}
        doc = generator._construir_docx(doc_data)
        texts = "\n".join(p.text for p in doc.paragraphs)
        assert "María García" not in texts

    def test_sin_notas_legales(self, generator):
        doc_data = {**_DOC_COMPLETO, "notas_legales": ""}
        doc = generator._construir_docx(doc_data)
        texts = "\n".join(p.text for p in doc.paragraphs)
        assert "⚠" not in texts

    def test_cuerpo_varios_parrafos(self, generator):
        cuerpo = ["Párrafo 1", "Párrafo 2", "Párrafo 3"]
        doc_data = {**_DOC_COMPLETO, "cuerpo": cuerpo}
        doc = generator._construir_docx(doc_data)
        texts = "\n".join(p.text for p in doc.paragraphs)
        for p in cuerpo:
            assert p in texts


# ── EscritoGenerator._generar_escrito_con_claude ────────────────────────────

class TestGenerarConClaude:
    """Cubre _generar_escrito_con_claude (lines 80-105)."""

    @patch("src.tramites.documentos.escrito.call_claude")
    def test_llamada_claude(self, mock_call_claude, generator):
        mock_call_claude.return_value = {"titulo": "Test"}
        datos = {"instrucciones": "carta formal", "nombre": "Ana"}
        result = generator._generar_escrito_con_claude("carta_presentacion", datos)
        assert result == {"titulo": "Test"}
        mock_call_claude.assert_called_once()

    @patch("src.tramites.documentos.escrito.call_claude")
    def test_prompt_incluye_tipo(self, mock_call_claude, generator):
        mock_call_claude.return_value = {}
        generator._generar_escrito_con_claude("queja_formal", {})
        prompt_arg = mock_call_claude.call_args[0][0]
        full_prompt = " ".join(m.get("content", "") for m in prompt_arg)
        assert "queja formal" in full_prompt

    @patch("src.tramites.documentos.escrito.call_claude")
    def test_tipo_desconocido(self, mock_call_claude, generator):
        """Tipo no registrado en DESCRIPCIONES → usa el tipo como descripción."""
        mock_call_claude.return_value = {}
        generator._generar_escrito_con_claude("tipo_inexistente", {})
        prompt_arg = mock_call_claude.call_args[0][0]
        full_prompt = " ".join(m.get("content", "") for m in prompt_arg)
        assert "tipo_inexistente" in full_prompt


# ── EscritoGenerator._preguntar ─────────────────────────────────────────────

class TestPreguntar:
    @patch("builtins.input", return_value="  valor  ")
    def test_strip(self, mock_input, generator):
        result = generator._preguntar("Nombre")
        assert result == "valor"

    @patch("builtins.input", return_value="texto")
    def test_valor_directo(self, mock_input, generator):
        result = generator._preguntar("Campo")
        assert result == "texto"


# ── EscritoGenerator.generar_interactivo ────────────────────────────────────

class TestGenerarInteractivo:
    """Cubre generar_interactivo (lines 211-266)."""

    @patch("builtins.input", return_value="")
    @patch("src.tramites.documentos.escrito.call_claude")
    def test_exitoso(self, mock_call_claude, mock_input, tmp_output):
        """Flujo completo exitoso con tipo default."""
        mock_call_claude.return_value = {"titulo": "Documento"}
        gen = EscritoGenerator()
        result = gen.generar_interactivo()
        assert result["status"] == "ok"
        assert ".docx" in result["archivo"]
        assert result["tipo"] == "escrito_libre"
        assert Path(result["archivo"]).exists()
        Path(result["archivo"]).unlink()

    @patch("builtins.input", return_value="")
    @patch("src.tramites.documentos.escrito.call_claude")
    def test_con_tipo_especifico(self, mock_call_claude, mock_input, tmp_output):
        """Selecciona tipo 3 (solicitud_empleo)."""
        mock_call_claude.return_value = {"titulo": "Solicitud"}
        gen = EscritoGenerator()
        with patch("builtins.input", side_effect=["3", "", "", "", ""]):
            result = gen.generar_interactivo()
        assert result["tipo"] == "solicitud_empleo"

    @patch("src.tramites.documentos.escrito.call_claude")
    def test_error_claude(self, mock_call_claude, tmp_output):
        """Line 244-245: call_claude falla → status error."""
        mock_call_claude.side_effect = Exception("API error")
        with patch("builtins.input", return_value=""):
            gen = EscritoGenerator()
            result = gen.generar_interactivo()
        assert result["status"] == "error"
        assert "API error" in result["error"]

    def test_error_construir_docx(self, tmp_output):
        """Error en _construir_docx → status error (con try/except)."""
        gen = EscritoGenerator()
        gen._construir_docx = MagicMock(side_effect=ValueError("doc error"))
        with (
            patch("builtins.input", return_value=""),
            patch("src.tramites.documentos.escrito.call_claude",
                  return_value={"titulo": ""}),
        ):
            result = gen.generar_interactivo()
        assert result["status"] == "error"
        assert "doc error" in result["error"]

    @patch("builtins.input", return_value="")
    @patch("src.tramites.documentos.escrito.call_claude")
    def test_nombre_archivo_incluye_tipo(self, mock_call_claude, mock_input,
                                         tmp_output):
        """El nombre del archivo incluye el tipo y timestamp."""
        mock_call_claude.return_value = {"titulo": ""}
        gen = EscritoGenerator()
        with patch("src.tramites.documentos.escrito.datetime") as mock_dt:
            mock_dt.now.return_value.strftime.return_value = "20260327_120000"
            result = gen.generar_interactivo()
        assert "escrito_libre-20260327_120000.docx" in result["archivo"]
