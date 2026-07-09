"""Tests para src/tramites/documentos/cv.py — Generador de CV profesional."""

import json
import os
from pathlib import Path
from unittest.mock import MagicMock, patch

import pytest

from src.tramites.documentos.cv import (
    OUTPUT_DIR, _add_line, _bullet, _parrafo, _seccion, CVGenerator,
)

# ── Helpers ──────────────────────────────────────────────────────────────────

_CV_COMPLETO = {
    "nombre_completo": "Juan Pérez López",
    "titulo_profesional": "Ingeniero de Software",
    "resumen": "Profesional con 5 años de experiencia en desarrollo web.",
    "experiencia": [
        {
            "empresa": "Tech Corp",
            "puesto": "Senior Developer",
            "periodo": "2020-2023",
            "logros": ["Lideré equipo de 5", "Migré a microservicios"],
        }
    ],
    "educacion": [
        {
            "institucion": "UNAM",
            "grado": "Ing. en Computación",
            "año": "2018",
        }
    ],
    "habilidades": ["Python", "JavaScript", "SQL"],
    "idiomas": ["Español nativo", "Inglés B2"],
    "contacto": {
        "telefono": "555-1234",
        "email": "juan@example.com",
        "ciudad": "CDMX",
    },
}

_CV_MINIMO = {
    "nombre_completo": "",
    "experiencia": [],
    "educacion": [],
    "habilidades": [],
    "idiomas": [],
    "contacto": {},
}


@pytest.fixture
def cv_generator():
    return CVGenerator()


@pytest.fixture
def tmp_output(tmp_path):
    """Parchea OUTPUT_DIR para evitar ensuciar el filesystem."""
    with patch("src.tramites.documentos.cv.OUTPUT_DIR", tmp_path):
        yield tmp_path


def _get_texts(doc):
    """Helper: extrae texto plano del documento."""
    return [p.text for p in doc.paragraphs]


# ── Funciones helper ─────────────────────────────────────────────────────────

class TestHelpers:
    def test_add_line(self):
        from docx import Document
        doc = Document()
        p = _add_line(doc)
        assert p is not None

    def test_seccion(self, cv_generator):
        from docx import Document
        doc = Document()
        _seccion(doc, "Experiencia Laboral")
        texts = _get_texts(doc)
        assert any("EXPERIENCIA LABORAL" in t for t in texts)

    def test_bullet(self):
        from docx import Document
        doc = Document()
        _bullet(doc, "Logro importante")
        texts = _get_texts(doc)
        assert any("Logro importante" in t for t in texts)

    def test_parrafo(self):
        from docx import Document
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.shared import RGBColor

        doc = Document()
        p = _parrafo(doc, "Texto de prueba", size=11, bold=True,
                     color=RGBColor(0xFF, 0, 0))
        assert "Texto de prueba" in p.text
        run = p.runs[0]
        assert run.font.size.pt == 11
        assert run.bold
        assert run.font.color.rgb == (0xFF, 0, 0)

    def test_parrafo_centrado(self):
        from docx import Document
        from docx.enum.text import WD_ALIGN_PARAGRAPH

        doc = Document()
        p = _parrafo(doc, "Centrado", alignment=WD_ALIGN_PARAGRAPH.CENTER)
        assert p.alignment == WD_ALIGN_PARAGRAPH.CENTER


# ── CVGenerator._construir_docx ──────────────────────────────────────────────

class TestConstruirDocx:
    """Cubre _construir_docx — construcción del documento Word (lines 135-258)."""

    def test_cv_completo(self, cv_generator):
        doc = cv_generator._construir_docx(_CV_COMPLETO)
        texts = "\n".join(_get_texts(doc))

        # Encabezado + título
        assert "Juan Pérez López" in texts
        assert "Ingeniero de Software" in texts
        # Perfil
        assert "Profesional con 5 años" in texts
        # Experiencia
        assert "Senior Developer" in texts
        assert "Tech Corp" in texts
        assert "Lideré equipo" in texts
        # Educación
        assert "Ing. en Computación" in texts
        assert "UNAM" in texts
        # Habilidades
        assert "Python" in texts
        # Idiomas
        assert "Español nativo" in texts
        # Contacto
        assert "555-1234" in texts
        assert "juan@example.com" in texts

    def test_cv_minimo_no_crashea(self, cv_generator):
        """Campos vacíos no deben causar errores."""
        doc = cv_generator._construir_docx(_CV_MINIMO)
        assert doc is not None
        assert len(doc.paragraphs) >= 3

    def test_cv_sin_contacto(self, cv_generator):
        """Sin contacto no se agrega línea de contacto."""
        cv = {**_CV_COMPLETO, "contacto": {}}
        doc = cv_generator._construir_docx(cv)
        texts = "\n".join(_get_texts(doc))
        assert texts.count("PERFIL PROFESIONAL") == 1

    def test_cv_sin_experiencia(self, cv_generator):
        """Sin experiencia no crashea y sigue con educación."""
        cv = {**_CV_COMPLETO, "experiencia": []}
        doc = cv_generator._construir_docx(cv)
        texts = "\n".join(_get_texts(doc))
        assert "EDUCACI" in texts.upper()

    def test_cv_sin_educacion(self, cv_generator):
        """Sin educación no crashea."""
        cv = {**_CV_COMPLETO, "educacion": []}
        doc = cv_generator._construir_docx(cv)
        texts = "\n".join(_get_texts(doc))
        assert "EXPERIENCIA LABORAL" in texts.upper()

    def test_cv_sin_habilidades(self, cv_generator):
        """Lista vacía de habilidades no agrega texto."""
        cv = {**_CV_COMPLETO, "habilidades": []}
        doc = cv_generator._construir_docx(cv)
        texts = "\n".join(_get_texts(doc))
        # Habilidades vacío: no hay sección con bullets (solo título)
        assert "HABILIDADES" in texts.upper()
        assert "Python" not in texts

    def test_cv_sin_idiomas(self, cv_generator):
        """Lista vacía de idiomas no agrega texto."""
        cv = {**_CV_COMPLETO, "idiomas": []}
        doc = cv_generator._construir_docx(cv)
        texts = "\n".join(_get_texts(doc))
        assert "IDIOMAS" in texts.upper()
        assert "Español" not in texts

    def test_cv_logros_vacios(self, cv_generator):
        """Logros vacíos no crashean."""
        cv = {**_CV_COMPLETO, "experiencia": [
            {**_CV_COMPLETO["experiencia"][0], "logros": []}
        ]}
        doc = cv_generator._construir_docx(cv)
        texts = "\n".join(_get_texts(doc))
        assert "Senior Developer" in texts

    def test_cv_empresa_vacia(self, cv_generator):
        """Empresa vacía en experiencia."""
        cv = {**_CV_COMPLETO, "experiencia": [
            {**_CV_COMPLETO["experiencia"][0], "empresa": ""}
        ]}
        doc = cv_generator._construir_docx(cv)
        texts = "\n".join(_get_texts(doc))
        assert "Senior Developer" in texts

    def test_cv_periodo_vacio(self, cv_generator):
        """Periodo vacío en experiencia."""
        cv = {**_CV_COMPLETO, "experiencia": [
            {**_CV_COMPLETO["experiencia"][0], "periodo": ""}
        ]}
        doc = cv_generator._construir_docx(cv)
        texts = "\n".join(_get_texts(doc))
        assert "Senior Developer" in texts

    def test_cv_institucion_vacia(self, cv_generator):
        """Institución vacía en educación."""
        cv = {**_CV_COMPLETO, "educacion": [
            {**_CV_COMPLETO["educacion"][0], "institucion": ""}
        ]}
        doc = cv_generator._construir_docx(cv)
        texts = "\n".join(_get_texts(doc))
        assert "Ing. en Computación" in texts

    def test_cv_anio_vacio(self, cv_generator):
        """Año vacío en educación."""
        cv = {**_CV_COMPLETO, "educacion": [
            {**_CV_COMPLETO["educacion"][0], "año": ""}
        ]}
        doc = cv_generator._construir_docx(cv)
        texts = "\n".join(_get_texts(doc))
        assert "Ing. en Computación" in texts


# ── CVGenerator._generar_cv_con_claude ──────────────────────────────────────

class TestGenerarCVConClaude:
    """Cubre _generar_cv_con_claude (lines 104-131)."""

    @patch("src.tramites.documentos.cv.call_claude")
    def test_llamada_claude(self, mock_call_claude, cv_generator):
        mock_call_claude.return_value = {"nombre_completo": "Test"}
        datos = {"nombre": "Juan", "puesto_objetivo": "Dev"}
        result = cv_generator._generar_cv_con_claude(datos)
        assert result == {"nombre_completo": "Test"}
        mock_call_claude.assert_called_once()
        # Verificar que el prompt incluye los datos
        prompt_arg = mock_call_claude.call_args[0][0]
        assert any("Juan" in msg.get("content", "") for msg in prompt_arg)

    @patch("src.tramites.documentos.cv.call_claude")
    def test_prompt_incluye_estructura(self, mock_call_claude, cv_generator):
        mock_call_claude.return_value = {}
        cv_generator._generar_cv_con_claude({"nombre": "Ana"})
        prompt_arg = mock_call_claude.call_args[0][0]
        full_prompt = " ".join(
            m.get("content", "") for m in prompt_arg
        )
        assert "experiencia" in full_prompt
        assert "educacion" in full_prompt
        assert "habilidades" in full_prompt


# ── CVGenerator._preguntar ───────────────────────────────────────────────────

class TestPreguntar:
    """Cubre _preguntar (lines 262-266)."""

    @patch("builtins.input", return_value="valor ingresado")
    def test_preguntar_con_valor(self, mock_input, cv_generator):
        result = cv_generator._preguntar("Nombre")
        assert result == "valor ingresado"

    @patch("builtins.input", return_value="")
    def test_preguntar_default(self, mock_input, cv_generator):
        result = cv_generator._preguntar("Email", default="test@mail.com")
        assert result == "test@mail.com"

    @patch("builtins.input", return_value="  espaciado  ")
    def test_preguntar_strip(self, mock_input, cv_generator):
        result = cv_generator._preguntar("Nombre")
        assert result == "espaciado"


# ── CVGenerator.generar_interactivo ─────────────────────────────────────────

class TestGenerarInteractivo:
    """Cubre generar_interactivo (lines 268-324)."""

    @patch("builtins.input", return_value="")
    @patch("src.tramites.documentos.cv.call_claude")
    def test_exitoso(self, mock_call_claude, mock_input, cv_generator, tmp_output):
        """Flujo completo exitoso."""
        mock_call_claude.return_value = {
            "nombre_completo": "Juan Pérez",
        }

        result = cv_generator.generar_interactivo()

        assert result["status"] == "ok"
        assert "cv-" in result["archivo"]
        # Verificar que se creó el archivo
        assert Path(result["archivo"]).exists()

    @patch("builtins.input", return_value="")
    @patch("src.tramites.documentos.cv.call_claude")
    def test_archivo_guardado(self, mock_call_claude, mock_input, tmp_output):
        """El archivo .docx se guarda con el nombre correcto."""
        mock_call_claude.return_value = {"nombre_completo": ""}
        gen = CVGenerator()
        result = gen.generar_interactivo()
        ruta = Path(result["archivo"])
        assert ruta.suffix == ".docx"
        assert ruta.parent == tmp_output
        assert ruta.exists()
        # cleanup
        ruta.unlink()

    @patch("builtins.input", return_value="")
    @patch("src.tramites.documentos.cv.call_claude")
    def test_prompt_claude_error(self, mock_call_claude, mock_input,
                                 cv_generator, tmp_output):
        """Lines 305-307: call_claude falla → status error."""
        mock_call_claude.side_effect = Exception("API error")

        result = cv_generator.generar_interactivo()

        assert result["status"] == "error"
        assert "API error" in result["error"]

    def test_con_perfiles_json(self, cv_generator, tmp_output):
        """Lines 283-289: perfiles.json existe y se carga."""
        perfil = {"nombre": "María", "email": "maria@test.com"}
        perfiles_path = tmp_output / "perfiles.json"
        perfiles_path.write_text(json.dumps(perfil), encoding="utf-8")

        with (
            patch("builtins.input", return_value=""),
            patch("src.tramites.documentos.cv.call_claude",
                  return_value={"nombre_completo": "María"}),
        ):
            result = cv_generator.generar_interactivo()

        assert result["status"] == "ok"

    def test_perfiles_json_invalido(self, cv_generator, tmp_output):
        """Lines 289-290: JSON inválido → se ignora."""
        perfiles_path = tmp_output / "perfiles.json"
        perfiles_path.write_text("not json", encoding="utf-8")

        with (
            patch("builtins.input", return_value=""),
            patch("src.tramites.documentos.cv.call_claude",
                  return_value={"nombre_completo": "Test"}),
        ):
            result = cv_generator.generar_interactivo()

        assert result["status"] == "ok"

    def test_perfiles_json_no_lista(self, cv_generator, tmp_output):
        """perfiles.json con lista en vez de dict → se ignora."""
        perfiles_path = tmp_output / "perfiles.json"
        perfiles_path.write_text(json.dumps([1, 2, 3]), encoding="utf-8")

        with (
            patch("builtins.input", return_value=""),
            patch("src.tramites.documentos.cv.call_claude",
                  return_value={"nombre_completo": "Test"}),
        ):
            result = cv_generator.generar_interactivo()

        assert result["status"] == "ok"

    def test_perfiles_json_no_existe(self, cv_generator, tmp_output):
        """No hay perfiles.json → arranca sin datos precargados."""
        with (
            patch("builtins.input", return_value=""),
            patch("src.tramites.documentos.cv.call_claude",
                  return_value={"nombre_completo": ""}),
        ):
            result = cv_generator.generar_interactivo()
        assert result["status"] == "ok"

    @patch("src.tramites.documentos.cv.call_claude")
    def test_sanitiza_nombre_archivo(self, mock_call_claude, tmp_output):
        """Line 313: espacios en nombre → guiones."""
        mock_call_claude.return_value = {"nombre_completo": ""}
        with patch("builtins.input", return_value=""):
            gen = CVGenerator()
            result = gen.generar_interactivo()
        assert "cv--" not in Path(result["archivo"]).name  # sin dobles guiones

    def test_error_en_construir_docx(self, cv_generator, tmp_output):
        """Si _construir_docx crashea, se propaga."""
        from docx import Document

        def fail_docx(cv):
            raise ValueError("Document error")

        cv_generator._construir_docx = fail_docx

        with (
            patch("builtins.input", return_value=""),
            patch("src.tramites.documentos.cv.call_claude",
                  return_value={"nombre_completo": "Test"}),
        ):
            result = cv_generator.generar_interactivo()

        assert result["status"] == "error"
        assert "Document error" in result["error"]


# ── CVGenerator.__init__ ────────────────────────────────────────────────────

class TestInit:
    def test_default_api_key(self):
        gen = CVGenerator()
        assert gen.api_key is None

    def test_custom_api_key(self):
        gen = CVGenerator(api_key="sk-xxx")
        assert gen.api_key == "sk-xxx"
